"""
Document text extraction service.

Provides extractors for PDF and DOCX files to enable AI-powered parsing
of meeting documents for activity creation.
"""

from abc import ABC, abstractmethod
from io import BytesIO
from typing import Protocol

from fastapi import UploadFile


class DocumentExtractor(Protocol):
    """Protocol for document text extractors."""

    async def extract_text(self, file: UploadFile) -> str:
        """Extract text content from a document file."""
        ...


class PDFExtractor:
    """Extract text from PDF files using pypdf."""

    async def extract_text(self, file: UploadFile) -> str:
        """Extract text content from a PDF file."""
        from pypdf import PdfReader

        content = await file.read()
        await file.seek(0)  # Reset file position for potential reuse

        pdf_file = BytesIO(content)
        reader = PdfReader(pdf_file)

        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        return "\n\n".join(text_parts)


class DOCXExtractor:
    """Extract text from Word documents using python-docx."""

    async def extract_text(self, file: UploadFile) -> str:
        """Extract text content from a DOCX file."""
        from docx import Document

        content = await file.read()
        await file.seek(0)  # Reset file position for potential reuse

        docx_file = BytesIO(content)
        doc = Document(docx_file)

        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n\n".join(text_parts)


SUPPORTED_CONTENT_TYPES = {
    "application/pdf": PDFExtractor,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DOCXExtractor,
}


def get_extractor(content_type: str) -> DocumentExtractor | None:
    """Get the appropriate extractor for a given content type."""
    extractor_class = SUPPORTED_CONTENT_TYPES.get(content_type)
    if extractor_class:
        return extractor_class()
    return None


def is_supported_content_type(content_type: str) -> bool:
    """Check if the content type is supported for document parsing."""
    return content_type in SUPPORTED_CONTENT_TYPES
